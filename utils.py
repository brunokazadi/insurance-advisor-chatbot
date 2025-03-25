import os
import time
import speech_recognition as sr
from dotenv import load_dotenv
from groq import Groq
from gtts import gTTS
import gradio as gr
from static import JS_ANIMATE, JS_THEME, STYLE, THEME_RESET_SCRIPT

# ----------------------------------------------------------------------------- 
# Configuration & Environment Setup 
# -----------------------------------------------------------------------------

# Load API key 
API_KEY = os.getenv("GROQ_API_KEY")

client = Groq(api_key=API_KEY)

# Currency configuration and theme definition
CURRENCY_MAP = {
    "USD": ("$", 50, 2000),
    "EUR": ("€", 50, 1800),
    "CAD": ("CA$", 50, 1600)
}
custom_theme = gr.themes.Default(primary_hue="blue", secondary_hue="blue")

# Translation dictionaries for UI elements
TRANSLATIONS = {
    "🇬🇧 English": {
        "chat_tab": "💬 Chat",
        "policy_tab": "🔍 Policy Finder",
        "insurance_title": "🏦 Insurance Advisor Chatbot",
        "insurance_subtitle": "Discuss your insurance needs and get personalized policy recommendations!",
        "policy_title": "🔍 Policy Finder",
        "policy_subtitle": "Enter your requirements to receive tailored insurance policy recommendations.",
        "placeholder_text": "Enter your question or upload voice or file to consult your insurance advisor...",
        "tts_label": "Enable Text-to-Speech",
        "examples_label": "Examples",
        "policy_details_label": "📝 Policy Details (Describe your needs)",
        "policy_details_placeholder": "E.g., I need comprehensive car insurance with roadside assistance",
        "insurance_type_label": "Insurance Type",
        "coverage_label": "Coverage Amount (optional)",
        "coverage_placeholder": "E.g., $50,000",
        "currency_label": "Select Currency",
        "budget_label": "Premium Budget ({}) (optional)",
        "people_label": "Number of Insured Individuals (optional)",
        "term_label": "Policy Term (optional)",
        "term_placeholder": "E.g., 1 year, 5 years",
        "generate_btn": "Generate Recommendation",
        "recommendation_label": "Recommendation",
        "generating_text": "**Generating policy recommendation...**",
        "examples": [
            "I need help finding a comprehensive car insurance policy.",
            "What are the benefits of term life insurance?",
            "Can you recommend a home insurance policy for a new homeowner?",
            "What should I consider for a health insurance plan?"
        ],
        "file_unsupported": "I've uploaded a file with unsupported format: '{}'. The insurance advisor can process text, document, and audio files. Please upload a supported file format or type your question.",
        "image_unsupported": "I've uploaded an image file named '{}'. Unfortunately, image analysis is not yet supported. Please describe what the image contains so I can assist you better.",
        "document_uploaded": "I've uploaded a document named '{}'. Please analyze this file for insurance-relevant information:",
        "audio_uploaded": "I've uploaded an audio file named '{}' with the following content:",
        "error_pdf_import": "Error: PyPDF2 library is not installed. Unable to read PDF file {}.",
        "error_docx_import": "Error: python-docx library is not installed. Unable to read {} file {}.",
        "error_reading": "Error reading {} file {}: {}"
    },
    "🇫🇷 Français": {
        "chat_tab": "💬 Discussion",
        "policy_tab": "🔍 Recherche de Police",
        "insurance_title": "🏦 Chatbot Conseiller en Assurance",
        "insurance_subtitle": "Discutez de vos besoins en assurance et obtenez des recommandations personnalisées!",
        "policy_title": "🔍 Recherche de Police",
        "policy_subtitle": "Entrez vos exigences pour recevoir des recommandations de polices d'assurance sur mesure.",
        "placeholder_text": "Saisissez votre question ou téléchargez un fichier vocal ou un document pour consulter votre conseiller en assurance...",
        "tts_label": "Activer la Synthèse Vocale",
        "examples_label": "Exemples",
        "policy_details_label": "📝 Détails de la Police (Décrivez vos besoins)",
        "policy_details_placeholder": "Ex: J'ai besoin d'une assurance auto complète avec assistance routière",
        "insurance_type_label": "Type d'Assurance",
        "coverage_label": "Montant de Couverture (optionnel)",
        "coverage_placeholder": "Ex: 50 000 €",
        "currency_label": "Sélectionnez la Devise",
        "budget_label": "Budget de Prime ({}) (optionnel)",
        "people_label": "Nombre de Personnes Assurées (optionnel)",
        "term_label": "Durée de la Police (optionnel)",
        "term_placeholder": "Ex: 1 an, 5 ans",
        "generate_btn": "Générer une Recommandation",
        "recommendation_label": "Recommandation",
        "generating_text": "**Génération de recommandation en cours...**",
        "examples": [
            "J'ai besoin d'aide pour trouver une police d'assurance automobile complète.",
            "Quels sont les avantages de l'assurance vie temporaire?",
            "Pouvez-vous recommander une police d'assurance habitation pour un nouveau propriétaire?",
            "Que dois-je considérer pour un régime d'assurance maladie?"
        ],
        "file_unsupported": "J'ai téléchargé un fichier avec un format non pris en charge: '{}'. Le conseiller en assurance peut traiter des fichiers texte, document et audio. Veuillez télécharger un format de fichier pris en charge ou tapez votre question.",
        "image_unsupported": "J'ai téléchargé un fichier image nommé '{}'. Malheureusement, l'analyse d'image n'est pas encore prise en charge. Veuillez décrire ce que contient l'image pour que je puisse mieux vous aider.",
        "document_uploaded": "J'ai téléchargé un document nommé '{}'. Veuillez analyser ce fichier pour des informations pertinentes sur l'assurance:",
        "audio_uploaded": "J'ai téléchargé un fichier audio nommé '{}' avec le contenu suivant:",
        "error_pdf_import": "Erreur: La bibliothèque PyPDF2 n'est pas installée. Impossible de lire le fichier PDF {}.",
        "error_docx_import": "Erreur: La bibliothèque python-docx n'est pas installée. Impossible de lire le fichier {} {}.",
        "error_reading": "Erreur de lecture du fichier {} {}: {}"
    }
}

# Insurance type options by language
INSURANCE_TYPES = {
    "🇬🇧 English": ["Auto", "Home", "Life", "Health", "Travel"],
    "🇫🇷 Français": ["Auto", "Habitation", "Vie", "Santé", "Voyage"]
}

# ----------------------------------------------------------------------------- 
# Business Logic Functions 
# -----------------------------------------------------------------------------

def transcribe_audio(audio_path):
    """Transcribe an audio file using Google's speech recognition."""
    recognizer = sr.Recognizer()
    with sr.AudioFile(audio_path) as source:
        audio = recognizer.record(source)
    try:
        return recognizer.recognize_google(audio)
    except sr.UnknownValueError:
        return "Cannot read voice"
    except sr.RequestError:
        return "Not available"

def read_file_content(file_path, language="🇬🇧 English"):
    """Read and extract text from PDF, DOC/DOCX, or TXT files with improved metadata."""
    translations = TRANSLATIONS[language]
    file_name = os.path.basename(file_path)
    ext = os.path.splitext(file_path)[1].lower()
    
    # Initialize response with metadata
    result = f"File Name: {file_name}\nFile Type: {ext[1:].upper()}\n\n"
    
    if ext == ".txt":
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
                result += "Content:\n" + content
                return result
        except Exception as e:
            return translations["error_reading"].format("TXT", file_name, str(e))
    elif ext == ".pdf":
        try:
            import PyPDF2
        except ImportError:
            return translations["error_pdf_import"].format(file_name)
        try:
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                result += f"Total Pages: {len(reader.pages)}\n\nContent:\n"
                
                # Extract and add page numbers for better context
                for i, page in enumerate(reader.pages):
                    extracted = page.extract_text()
                    if extracted:
                        result += f"\n--- Page {i+1} ---\n{extracted}\n"
                return result
        except Exception as e:
            return translations["error_reading"].format("PDF", file_name, str(e))
    elif ext in [".doc", ".docx"]:
        try:
            import docx
        except ImportError:
            return translations["error_docx_import"].format(ext, file_name)
        try:
            doc = docx.Document(file_path)
            result += f"Total Paragraphs: {len(doc.paragraphs)}\n\nContent:\n"
            text = "\n".join([para.text for para in doc.paragraphs if para.text])
            result += text
            return result
        except Exception as e:
            return translations["error_reading"].format(ext, file_name, str(e))
    else:
        return f"Unsupported file format: {ext}. The insurance advisor can process .txt, .pdf, .doc and .docx files."

def process_input(history, message, language="🇬🇧 English"):
    """Process user message and uploaded files with enhanced context management."""
    if history is None:
        history = []
    
    translations = TRANSLATIONS[language]
    user_text = ""
    files = message.get("files", [])
    if files:
        for file in files:
            file_name = os.path.basename(file)
            
            # Process voice files
            if file.endswith(".wav") or file.endswith(".mp3"):
                transcribed_text = transcribe_audio(file)
                user_text = translations["audio_uploaded"].format(file_name) + "\n\n" + transcribed_text
                history.append((user_text, ""))
            
            # Process document files (PDF, DOC/DOCX, TXT)
            elif (file.endswith(".pdf") or 
                  file.endswith(".doc") or 
                  file.endswith(".docx") or 
                  file.endswith(".txt")):
                file_text = read_file_content(file, language)
                user_text = translations["document_uploaded"].format(file_name) + "\n\n" + file_text
                history.append((user_text, ""))
            
            # Handle image files with placeholder (for future implementation)
            elif (file.endswith(".jpg") or 
                  file.endswith(".jpeg") or 
                  file.endswith(".png") or
                  file.endswith(".gif")):
                user_text = translations["image_unsupported"].format(file_name)
                history.append((user_text, ""))
            
            else:
                user_text = translations["file_unsupported"].format(file_name)
                history.append((user_text, ""))
    
    if message.get("text"):
        user_text = message["text"]
        history.append((user_text, ""))
    
    # Return updated history and a fresh multimodal textbox for the next input
    return history, gr.MultimodalTextbox(value=None, interactive=True)

def generate_policy_recommendation(policy_details, insurance_type, coverage, budget, policy_term, num_people, currency, language):
    """Generate an insurance policy recommendation using the chatbot backend."""
    prompt_parts = ["Generate an insurance policy recommendation."]
    if insurance_type and insurance_type.strip():
        prompt_parts.append(f"Insurance Type: {insurance_type}.")
    if coverage and coverage.strip():
        prompt_parts.append(f"Coverage Amount: {coverage}.")
    if policy_details and policy_details.strip():
        prompt_parts.append(f"Policy Details: {policy_details}.")
    if budget:
        prompt_parts.append(f"Budget: {budget} {currency}.")
    if policy_term and policy_term.strip():
        prompt_parts.append(f"Policy Term: {policy_term}.")
    if num_people and num_people > 1:
        prompt_parts.append(f"Number of Insured Individuals: {num_people}.")
    
    user_prompt = " ".join(prompt_parts)
    messages = [
        {"role": "system", "content": (
            "Your name is Harvey Specter. You are an expert insurance advisor who ONLY provides information about insurance. "
            "Provide a structured recommendation for an insurance policy based on the user's requirements. "
            "Include details such as policy benefits, potential risks, and any clarifying questions that might help "
            "the client understand the policy. If asked about non-insurance topics, politely decline and "
            "redirect the conversation to insurance matters. "
            f"IMPORTANT: You must ALWAYS respond ONLY in {language} regardless of the language used in the user's input. "
            f"Even if the user asks you in a different language, you must respond only in {language}."
        )},
        {"role": "user", "content": user_prompt}
    ]
    
    try:
        completion = client.chat.completions.create(
            model="llama3-70b-8192",
            messages=messages,
            temperature=0.7,
            max_tokens=2048,
            top_p=0.9,
        )
        recommendation_text = completion.choices[0].message.content
        return recommendation_text
    except Exception as e:
        if "Français" in language:
            return f"**Erreur lors de la génération de la recommandation: {str(e)[:100]}... Veuillez réessayer.**"
        else:
            return f"**Error generating recommendation: {str(e)[:100]}... Please try again.**"

def chat_with_bot_stream(user_input, audio, language, history):
    """Stream responses from the chatbot with improved file handling."""
    if history is None:
        history = []
    history, _ = process_input(history, user_input, language)
    
    system_prompt = (
        "Your name is Harvey Specter. You are an expert insurance advisor with over 10 years of experience. "
        "You are ONLY authorized to answer questions related to insurance topics. "
        "You help clients find insurance policies that best match their needs, including auto, home, life, or health insurance. "
        
        "When the user uploads documents or files, you should:\n"
        "1. Analyze the content for insurance-relevant information\n"
        "2. Extract key details like policy numbers, coverage amounts, premium details, and conditions\n"
        "3. Provide a concise summary of the document's relevance to insurance\n"
        "4. Answer questions about the document in the context of insurance advice\n"
        "5. If the document isn't insurance-related, politely explain this and ask if they have insurance questions\n"
        
        "If a user asks a question that is not related to insurance, politely inform them that you can only "
        "assist with insurance-related inquiries and suggest they ask about insurance topics instead. "
        "Provide clear, detailed, and personalized advice for insurance topics only. "
        f"IMPORTANT: You must ALWAYS respond ONLY in {language} regardless of the language used in the user's input. "
        f"Even if the user asks you in a different language, you must respond only in {language}."
    )
    
    # Include more context by using up to 5 turns of conversation instead of 3
    messages = [{"role": "system", "content": system_prompt}]
    for user_msg, ai_response in history[-5:]:  # Increased from -3 to -5 for better context
        if user_msg:
            messages.append({"role": "user", "content": user_msg})
        if ai_response:
            if isinstance(ai_response, tuple):
                messages.append({"role": "assistant", "content": ai_response[0]})
            else:
                messages.append({"role": "assistant", "content": ai_response})
    
    if history[-1][1] == "":
        messages.append({"role": "user", "content": history[-1][0]})
    
    try:
        completion = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=messages,
            temperature=0.7,
            max_completion_tokens=5000,
            top_p=0.9,
            stream=True
        )
    except Exception as e:
        error_message = f"Error: {str(e)[:100]}..."
        offline_message = f"Advisor is currently offline, please wait a moment. {error_message if 'Français' not in language else ''}"
        if "Français" in language:
            offline_message = f"Le conseiller est actuellement hors ligne, veuillez patienter un moment. {error_message if 'Français' in language else ''}"
        yield [(offline_message, "")]
        return
    
    # Process response stream
    full_response = ""
    for chunk in completion:
        content = chunk.choices[0].delta.content or ""
        full_response += content  
        history[-1] = (history[-1][0], full_response)
        yield [(u, a) for u, a in history if u != "system"]
    
    if audio:
        try:
            tts_generating = "🔊 Generating text-to-speech..."
            if "Français" in language:
                tts_generating = "🔊 Génération de la synthèse vocale..."
            history.append(("", tts_generating))
            yield history
            
            language_map = {
                "🇬🇧 English": "en",
                "🇫🇷 Français": "fr"
            }
            tts_language = language_map.get(language, "en")
            audio_filename = f"bot_response_{int(time.time())}.mp3"
            tts = gTTS(full_response, lang=tts_language)
            tts.save(audio_filename)
            history.append(("", (audio_filename,)))
        except Exception as e:
            print(f"TTS failed: {e}")
    yield history

# ----------------------------------------------------------------------------- 
# UI Helper Functions
# -----------------------------------------------------------------------------

def update_budget_slider(currency, language):
    """Update the budget slider settings based on the selected currency."""
    symbol, min_val, max_val = CURRENCY_MAP[currency]
    lang_code = "🇬🇧 English" if "English" in language else "🇫🇷 Français"
    label_template = TRANSLATIONS[lang_code]["budget_label"]
    return gr.update(label=label_template.format(symbol), minimum=min_val, maximum=max_val)

def update_ui_language(language):
    """Update UI elements when language changes."""
    translations = TRANSLATIONS[language]
    
    # HTML updates
    insurance_title_html = f"""
        <h1>{translations["insurance_title"]}</h1>
        <h3 class="subtitle">{translations["insurance_subtitle"]}</h3>
    """
    policy_title_html = f"""
        <h1>{translations["policy_title"]}</h1>
        <h3 class="subtitle">{translations["policy_subtitle"]}</h3>
    """
    
    # Updates for components
    chat_tab_update = gr.update(label=translations["chat_tab"])
    policy_tab_update = gr.update(label=translations["policy_tab"])
    user_input_update = gr.update(placeholder=translations["placeholder_text"])
    audio_button_update = gr.update(label=translations["tts_label"])
    
    policy_details_update = gr.update(
        label=translations["policy_details_label"],
        placeholder=translations["policy_details_placeholder"]
    )
    insurance_type_update = gr.update(
        label=translations["insurance_type_label"],
        choices=INSURANCE_TYPES[language],
        value=INSURANCE_TYPES[language][0]
    )
    coverage_update = gr.update(
        label=translations["coverage_label"],
        placeholder=translations["coverage_placeholder"]
    )
    currency_update = gr.update(label=translations["currency_label"])
    
    # Get the current currency
    symbol, min_val, max_val = CURRENCY_MAP["USD"]  # Default value
    budget_update = gr.update(
        label=translations["budget_label"].format(symbol),
        minimum=min_val,
        maximum=max_val
    )
    
    people_update = gr.update(label=translations["people_label"])
    term_update = gr.update(
        label=translations["term_label"],
        placeholder=translations["term_placeholder"]
    )
    generate_btn_update = gr.update(value=translations["generate_btn"])
    recommendation_update = gr.update(label=translations["recommendation_label"])
    
    # Update example button visibility
    if language == "🇬🇧 English":
        english_examples_update = gr.update(visible=True)
        french_examples_update = gr.update(visible=False)
    else:
        english_examples_update = gr.update(visible=False)
        french_examples_update = gr.update(visible=True)
    
    # Update examples header
    examples_header_update = translations["examples_label"]
    
    return (
        chat_tab_update,
        policy_tab_update,
        insurance_title_html,
        policy_title_html,
        user_input_update,
        audio_button_update,
        policy_details_update,
        insurance_type_update,
        coverage_update,
        currency_update,
        budget_update,
        people_update,
        term_update,
        generate_btn_update,
        recommendation_update,
        english_examples_update,
        french_examples_update,
        examples_header_update
    )

# Custom button to insert example text
def use_example(example, textbox):
    """Insert example text into the textbox"""
    return example